import streamlit as st
import pandas as pd
import os
from dotenv import load_dotenv
import google.generativeai as genai
import io
from docx import Document

# Load environment variables from .env file
load_dotenv()

# Configure Gemini API key
genai.configure(api_key=os.getenv("GEMINI_API_KEY"))

# Function to generate interview questions using Gemini
def generate_interview_questions(domain, experience_level, complexity):
    """Generate interview questions and answers using Gemini."""
    try:
        # Define generation configuration
        generation_config = {
            "temperature": 1,
            "top_p": 0.95,
            "top_k": 40,
            "max_output_tokens": 8192,
            "response_mime_type": "text/plain",
        }

        # Initialize the Gemini model
        model = genai.GenerativeModel(
            model_name="gemini-1.5-flash",
            generation_config=generation_config,
        )

        # Start a chat session
        chat_session = model.start_chat()

        # Build the prompt dynamically based on the input domain
        prompt = (
            f"Generate a set of {complexity} interview questions and answers for a {domain} professional "
            f"with {experience_level} years of experience. The format should be clear and structured: "
            f"each question followed by the corresponding answer."
        )

        # Send the prompt and get the response
        response = chat_session.send_message(prompt)
        return response.text.strip()
    except Exception as e:
        st.error(f"Error fetching questions: {e}")
        return ""

# Function to export data to an Excel file and return the file as a BytesIO object
def export_to_excel(data):
    """Exports the data to an Excel file and returns the file in memory as a BytesIO object."""
    try:
        # Create a pandas DataFrame from the list of dictionaries
        df = pd.DataFrame(data)

        # Save to a BytesIO buffer using openpyxl (default engine for pandas)
        output = io.BytesIO()
        with pd.ExcelWriter(output, engine='openpyxl') as writer:
            df.to_excel(writer, index=False, sheet_name='Interview Questions')

        output.seek(0)  # Rewind the buffer
        return output
    except Exception as e:
        st.error(f"Error exporting to Excel: {e}")
        return None

# Function to export data to a Word document and return the file as a BytesIO object
def export_to_word(data):
    """Exports the data to a Word document and returns the file in memory as a BytesIO object."""
    try:
        # Create a Word document
        doc = Document()
        doc.add_heading('Interview Questions and Answers', 0)

        for qa in data:
            # Add question and answer to Word
            doc.add_heading(qa['Question'], level=1)
            doc.add_paragraph(qa['Answer'])

        # Save to a BytesIO buffer
        output = io.BytesIO()
        doc.save(output)
        output.seek(0)  # Rewind the buffer
        return output
    except Exception as e:
        st.error(f"Error exporting to Word: {e}")
        return None

# Main Streamlit app
def main():
    # Custom button styles
    custom_button_style = """
    <style>
    /* Style for Generate Questions button (Green with Black text) */
    .generate-button {
        background-color: #28a745; /* Green */
        color: black;
        border: none;
        padding: 15px 32px;
        text-align: center;
        text-decoration: none;
        display: inline-block;
        font-size: 16px;
        cursor: pointer;
        border-radius: 8px;
        transition: background 0.5s ease;
    }
    .generate-button:hover {
        background-color: #218838; /* Dark Green on hover */
    }

    /* Style for Export to Excel button (Light Red) */
    .export-excel-button {
        background-color: #f08080; /* Light Red */
        color: black;
        border: none;
        padding: 15px 32px;
        text-align: center;
        text-decoration: none;
        display: inline-block;
        font-size: 16px;
        cursor: pointer;
        border-radius: 8px;
        transition: background 0.5s ease;
    }
    .export-excel-button:hover {
        background-color: #e57373; /* Slightly darker Red on hover */
    }

    /* Style for Export to Word button (Light Blue) */
    .export-word-button {
        background-color: #add8e6; /* Light Blue */
        color: black;
        border: none;
        padding: 15px 32px;
        text-align: center;
        text-decoration: none;
        display: inline-block;
        font-size: 16px;
        cursor: pointer;
        border-radius: 8px;
        transition: background 0.5s ease;
    }
    .export-word-button:hover {
        background-color: #87ceeb; /* Slightly darker Blue on hover */
    }
    </style>
    """

    # Display the image at the top (centered)
    st.markdown("""
        <style>
            .top-center-image {
                display: flex;
                justify-content: center;
                align-items: center;
                margin-bottom: 20px;
            }
        </style>
        <div class="top-center-image">
            <img src="https://mazobeam.com/wp-content/uploads/2023/12/mazoid-1.png" alt="MazoBot" width="200"/>
        </div>
    """, unsafe_allow_html=True)

    # Display the title "MazoMate" after the image
    st.title("MazoMate")

    # Sidebar header for configuration
    st.sidebar.header("Configuration")

    # Text input for domain
    domain = st.sidebar.text_input("Enter a Programming/Area (e.g., Python, Java, C++, HR, Data Science, Marketing)")

    # Experience level slider
    experience_level = st.sidebar.slider("Experience Level (years)", min_value=1, max_value=50, step=1)

    # Complexity level radio buttons
    complexity = st.sidebar.radio("Question Complexity", ["Basic", "Intermediate", "Advanced"])

    st.write("### Selected Configuration")
    st.write(f"**Domain/Area**: {domain}")
    st.write(f"**Experience Level**: {experience_level} years")
    st.write(f"**Complexity**: {complexity}")

    # Inject the custom button styles
    st.markdown(custom_button_style, unsafe_allow_html=True)

    # Custom "Generate Questions" button with green background and black text
    if st.markdown('<button class="generate-button">Generate Questions</button>', unsafe_allow_html=True):
        st.info("Generating interview questions. Please wait...")

        # Generate questions
        generated_content = generate_interview_questions(domain, experience_level, complexity)

        if generated_content:
            st.success("Questions generated successfully!")
            st.write("### Generated Questions and Answers")
            st.text_area("Questions & Answers", generated_content, height=300)

            # Assuming the content comes back in a format where each Q&A pair is separated by newlines
            # Parsing the content into question-answer pairs
            qa_pairs = []
            lines = generated_content.split('\n')
            for i in range(0, len(lines), 2):  # Assuming questions and answers alternate
                question = lines[i].strip() if i < len(lines) else ""
                answer = lines[i+1].strip() if (i + 1) < len(lines) else ""
                qa_pairs.append({"Question": question, "Answer": answer})

            # Export options
            if qa_pairs:
                # Export to Excel button with light red color
                if st.markdown('<button class="export-excel-button">Export to Excel</button>', unsafe_allow_html=True):
                    excel_file = export_to_excel(qa_pairs)
                    if excel_file:
                        st.success("Ready to download Excel file!")
                        st.download_button(
                            label="Download Interview Questions Excel",
                            data=excel_file,
                            file_name="Mazo_Interview_Question_Answer.xlsx",
                            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                        )

                # Export to Word button with light blue color
                if st.markdown('<button class="export-word-button">Export to Word</button>', unsafe_allow_html=True):
                    word_file = export_to_word(qa_pairs)
                    if word_file:
                        st.success("Ready to download Word file!")
                        st.download_button(
                            label="Download Interview Questions Word",
                            data=word_file,
                            file_name="Mazo_Interview_Question_Answer.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )

if __name__ == "__main__":
    main()
